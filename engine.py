"""
KR Chart Rater - 엔진 (차트 분석 백엔드)
종목 차트 데이터 수집, 캔들차트 생성, LLM 비전 분석, 테마 ETF 탐색
"""

import os
import sys
import json
import time
import base64
import logging
import re
from datetime import datetime, timezone, timedelta
from pathlib import Path

import pandas as pd

# ============================================================
# 경로 설정
# ============================================================
if getattr(sys, "frozen", False):
    BASE_DIR = Path(sys.executable).parent.resolve()
else:
    BASE_DIR = Path(__file__).parent.resolve()

SECRETS_DIR = BASE_DIR / "secrets"
CHARTS_DIR = BASE_DIR / "charts"
RESULTS_DIR = BASE_DIR / "results"
LOGS_DIR = BASE_DIR / "logs"

for d in [CHARTS_DIR, RESULTS_DIR, LOGS_DIR]:
    d.mkdir(exist_ok=True)

KST = timezone(timedelta(hours=9))


# ============================================================
# 설정 로드
# ============================================================
def read_secret(filename):
    """secrets/ 폴더에서 API 키 등 비밀값 읽기. 환경변수 우선 (GitHub Actions 호환)."""
    # 환경변수 우선: "gemini_api_key.txt" → "GEMINI_API_KEY"
    env_key = filename.replace(".txt", "").upper()
    env_val = os.environ.get(env_key)
    if env_val:
        return env_val.strip()
    # 파일 방식
    path = SECRETS_DIR / filename
    if not path.exists():
        raise FileNotFoundError(f"환경변수 {env_key} 또는 파일 {path}를 찾을 수 없습니다")
    return path.read_text(encoding="utf-8").strip()


def load_config():
    """config.txt에서 key=value 설정 파싱"""
    config = {}
    config_path = BASE_DIR / "config.txt"
    if not config_path.exists():
        return config
    for line in config_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if "=" in line and not line.startswith("#"):
            key, value = line.split("=", 1)
            config[key.strip()] = value.strip()
    return config


CONFIG = load_config()

# 설정값 추출
LLM_PROVIDER = CONFIG.get("LLM_PROVIDER", "claude").lower()
CLAUDE_MODEL = CONFIG.get("CLAUDE_MODEL", "claude-sonnet-4-6")
GEMINI_MODEL = CONFIG.get("GEMINI_MODEL", "gemini-2.5-flash")
CHART_PERIOD = CONFIG.get("CHART_PERIOD", "1y")
CHART_MA_LINES = [int(x) for x in CONFIG.get("CHART_MA_LINES", "5,20,60,120").split(",")]
ETF_TOP_N = int(CONFIG.get("ETF_TOP_N_HOLDINGS", "10"))
SAVE_CHARTS = CONFIG.get("SAVE_CHARTS", "true").lower() == "true"

# 토큰 비용 단가 (USD per 1M tokens)
TOKEN_COST_RATES = {
    "claude-sonnet-4-6": {"input": 3.0, "output": 15.0},
    "claude-opus-4-6": {"input": 15.0, "output": 75.0},
    "gemini-2.5-flash": {"input": 0.15, "output": 0.60},
    "gemini-2.5-pro": {"input": 1.25, "output": 10.0},
    "gemini-3-flash-preview": {"input": 0.50, "output": 3.0},
    "gemini-3.1-pro-preview": {"input": 2.0, "output": 12.0},
    "gemini-3.1-flash-lite-preview": {"input": 0.25, "output": 1.5},
}


def _make_usage(provider, model, input_tokens, output_tokens):
    """토큰 사용량 딕셔너리 생성 + 비용 계산"""
    rates = TOKEN_COST_RATES.get(model, {"input": 0, "output": 0})
    input_cost = input_tokens * rates["input"] / 1_000_000
    output_cost = output_tokens * rates["output"] / 1_000_000
    return {
        "provider": provider,
        "model": model,
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "total_tokens": input_tokens + output_tokens,
        "input_cost_usd": input_cost,
        "output_cost_usd": output_cost,
        "total_cost_usd": input_cost + output_cost,
    }


def _accumulate_usage(accumulated, new_usage):
    """토큰 사용량 누적"""
    accumulated["input_tokens"] += new_usage["input_tokens"]
    accumulated["output_tokens"] += new_usage["output_tokens"]
    accumulated["total_tokens"] += new_usage["total_tokens"]
    accumulated["input_cost_usd"] += new_usage["input_cost_usd"]
    accumulated["output_cost_usd"] += new_usage["output_cost_usd"]
    accumulated["total_cost_usd"] += new_usage["total_cost_usd"]
    accumulated["api_calls"] += 1


# ============================================================
# 로거 설정
# ============================================================
def setup_logger(name="kr_chart_rater"):
    """파일 + 콘솔 이중 로깅 설정"""
    logger = logging.getLogger(name)
    if logger.handlers:
        return logger
    logger.setLevel(logging.DEBUG)

    # 파일 핸들러
    log_file = LOGS_DIR / f"{name}_{datetime.now(KST).strftime('%Y%m%d')}.log"
    fh = logging.FileHandler(str(log_file), encoding="utf-8")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(message)s"))

    # 콘솔 핸들러
    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(logging.Formatter("[%(levelname)s] %(message)s"))

    logger.addHandler(fh)
    logger.addHandler(ch)
    return logger


logger = setup_logger()


# ============================================================
# 종목 티커 매핑 (한글명 → 종목코드)
# ============================================================
_TICKER_MAP = None  # {한글명: (종목코드, 시장)} 캐시

# 주요 종목 내장 매핑 (pykrx 실패 시 폴백)
_BUILTIN_TICKERS = {
    "삼성전자": ("005930", "KOSPI"), "SK하이닉스": ("000660", "KOSPI"),
    "LG에너지솔루션": ("373220", "KOSPI"), "삼성바이오로직스": ("207940", "KOSPI"),
    "현대차": ("005380", "KOSPI"), "기아": ("000270", "KOSPI"),
    "셀트리온": ("068270", "KOSPI"), "KB금융": ("105560", "KOSPI"),
    "POSCO홀딩스": ("005490", "KOSPI"), "신한지주": ("055550", "KOSPI"),
    "NAVER": ("035420", "KOSPI"), "카카오": ("035720", "KOSPI"),
    "삼성SDI": ("006400", "KOSPI"), "LG화학": ("051910", "KOSPI"),
    "현대모비스": ("012330", "KOSPI"), "삼성물산": ("028260", "KOSPI"),
    "SK이노베이션": ("096770", "KOSPI"), "LG전자": ("066570", "KOSPI"),
    "삼성전기": ("009150", "KOSPI"), "SK텔레콤": ("017670", "KOSPI"),
    "한국전력": ("015760", "KOSPI"), "KT&G": ("033780", "KOSPI"),
    "하나금융지주": ("086790", "KOSPI"), "우리금융지주": ("316140", "KOSPI"),
    "HD현대중공업": ("329180", "KOSPI"), "한화에어로스페이스": ("012450", "KOSPI"),
    "두산에너빌리티": ("034020", "KOSPI"), "한화오션": ("042660", "KOSPI"),
    "HD한국조선해양": ("009540", "KOSPI"), "SK스퀘어": ("402340", "KOSPI"),
    "포스코퓨처엠": ("003670", "KOSPI"), "에코프로비엠": ("247540", "KOSDAQ"),
    "에코프로": ("086520", "KOSDAQ"), "엘앤에프": ("066970", "KOSDAQ"),
    "HLB": ("028300", "KOSDAQ"), "알테오젠": ("196170", "KOSDAQ"),
    "리가켐바이오": ("141080", "KOSDAQ"), "레인보우로보틱스": ("277810", "KOSDAQ"),
    "두산로보틱스": ("454910", "KOSPI"), "한미반도체": ("042700", "KOSDAQ"),
    "이수페타시스": ("007660", "KOSPI"), "LS ELECTRIC": ("010120", "KOSPI"),
    "HD현대일렉트릭": ("267260", "KOSPI"), "효성중공업": ("298040", "KOSPI"),
    "KT": ("030200", "KOSPI"), "LG": ("003550", "KOSPI"),
    "SK": ("034730", "KOSPI"), "한화솔루션": ("009830", "KOSPI"),
    "삼성생명": ("032830", "KOSPI"), "삼성화재": ("000810", "KOSPI"),
    "크래프톤": ("259960", "KOSPI"), "엔씨소프트": ("036570", "KOSPI"),
    "넷마블": ("251270", "KOSPI"), "펄어비스": ("263750", "KOSPI"),
    "카카오게임즈": ("293490", "KOSDAQ"), "위메이드": ("112040", "KOSDAQ"),
    "컴투스": ("078340", "KOSDAQ"), "하이브": ("352820", "KOSPI"),
    "JYP Ent.": ("035900", "KOSPI"), "SM": ("041510", "KOSPI"),
    "와이지엔터테인먼트": ("122870", "KOSPI"), "CJ ENM": ("035760", "KOSPI"),
    "한화시스템": ("272210", "KOSPI"), "LIG넥스원": ("079550", "KOSPI"),
    "현대로템": ("064350", "KOSPI"), "한국항공우주": ("047810", "KOSPI"),
    "대한항공": ("003490", "KOSPI"), "한전KPS": ("051600", "KOSPI"),
    "에이피알": ("278470", "KOSDAQ"), "클래시스": ("214150", "KOSDAQ"),
    "삼성중공업": ("010140", "KOSPI"), "HMM": ("011200", "KOSPI"),
}


TICKER_CACHE_PATH = BASE_DIR / "ticker_cache.json"
THEME_CACHE_PATH = BASE_DIR / "theme_cache.json"
WATCHLIST_PATH = BASE_DIR / "watchlist.json"


def _load_ticker_cache():
    """ticker_cache.json에서 종목 매핑 로드"""
    if not TICKER_CACHE_PATH.exists():
        return {}
    try:
        data = json.loads(TICKER_CACHE_PATH.read_text(encoding="utf-8"))
        # {name: [code, market]} → {name: (code, market)}
        return {name: tuple(val) for name, val in data.items()}
    except Exception as e:
        logger.warning(f"ticker_cache.json 로드 실패: {e}")
        return {}


def _save_ticker_cache(ticker_map):
    """종목 매핑을 ticker_cache.json에 저장"""
    try:
        data = {name: list(val) for name, val in ticker_map.items()}
        with open(TICKER_CACHE_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=1)
        logger.info(f"ticker_cache.json 저장: {len(data)}개 종목")
    except Exception as e:
        logger.warning(f"ticker_cache.json 저장 실패: {e}")


def _build_ticker_map():
    """종목 매핑 빌드: 내장 → 캐시 파일 → pykrx 순으로 로드"""
    global _TICKER_MAP
    if _TICKER_MAP is not None:
        return _TICKER_MAP

    # 1. 내장 매핑으로 초기화
    _TICKER_MAP = dict(_BUILTIN_TICKERS)

    # 2. 캐시 파일 로드 (내장보다 우선)
    cached = _load_ticker_cache()
    if cached:
        _TICKER_MAP.update(cached)
        logger.info(f"ticker_cache.json 로드: {len(cached)}개 종목")

    # 3. pykrx로 전체 목록 시도 (성공 시 캐시 자동 갱신)
    pykrx_loaded = False
    try:
        from pykrx import stock
        for days_ago in range(0, 30):
            date = (datetime.now(KST) - timedelta(days=days_ago)).strftime("%Y%m%d")
            try:
                kospi = stock.get_market_ticker_list(date, market="KOSPI")
                if kospi:
                    for code in kospi:
                        name = stock.get_market_ticker_name(code)
                        _TICKER_MAP[name] = (code, "KOSPI")
                    kosdaq = stock.get_market_ticker_list(date, market="KOSDAQ")
                    for code in kosdaq:
                        name = stock.get_market_ticker_name(code)
                        _TICKER_MAP[name] = (code, "KOSDAQ")
                    pykrx_loaded = True
                    logger.info(f"pykrx 종목 매핑 로드 성공 ({date}): KOSPI {len(kospi)}, KOSDAQ {len(kosdaq)}")
                    _save_ticker_cache(_TICKER_MAP)
                    break
            except Exception:
                continue
        if not pykrx_loaded:
            logger.warning("pykrx 종목 목록 로드 실패, 캐시/내장 매핑 사용")
    except ImportError:
        logger.warning("pykrx 미설치, 캐시/내장 매핑 사용")
    except Exception as e:
        logger.warning(f"pykrx 종목 목록 로드 실패: {e}, 캐시/내장 매핑 사용")

    logger.info(f"종목 매핑 빌드 완료: {len(_TICKER_MAP)}개 종목")
    return _TICKER_MAP


def refresh_ticker_cache():
    """
    KIND API에서 전체 종목 목록을 가져와 ticker_cache.json 갱신.
    pykrx가 안 되는 환경에서도 독립적으로 캐시 갱신 가능.
    Returns: 갱신된 종목 수
    """
    import requests

    headers = {"User-Agent": "Mozilla/5.0"}
    cache = dict(_BUILTIN_TICKERS)

    for market_type, market_name in [("stockMkt", "KOSPI"), ("kosdaqMkt", "KOSDAQ")]:
        url = "https://kind.krx.co.kr/corpgeneral/corpList.do"
        params = {"method": "download", "marketType": market_type}

        try:
            resp = requests.get(url, params=params, headers=headers, timeout=15)
            content = resp.content.decode("euc-kr", errors="replace")

            import pandas as pd
            from io import StringIO
            dfs = pd.read_html(StringIO(content))
            if dfs:
                df = dfs[0]
                count = 0
                for _, row in df.iterrows():
                    name = str(row.iloc[0]).strip()
                    code_raw = row.iloc[2]
                    if pd.isna(code_raw) or name == "nan":
                        continue
                    code = str(int(code_raw)).zfill(6) if isinstance(code_raw, (int, float)) else str(code_raw).strip()
                    if len(code) == 6 and code.isdigit():
                        cache[name] = (code, market_name)
                        count += 1
                logger.info(f"KIND {market_name}: {count}개 종목 로드")
        except Exception as e:
            logger.error(f"KIND {market_name} 로드 실패: {e}")

    _save_ticker_cache(cache)

    # 현재 맵도 갱신
    global _TICKER_MAP
    if _TICKER_MAP is not None:
        _TICKER_MAP.update(cache)

    return len(cache)


# ============================================================
# 워치리스트 (watchlist.json) — 다중 리스트 지원 (v2)
# ============================================================
_DEFAULT_LIST_CONFIG = {
    "provider": None,
    "email_to": None,
    "notion_watchlist_db": None,
    "notion_report_db": None,
}


def _migrate_watchlist_v1_to_v2(data):
    """v1 {stocks: [...]} → v2 {version, active_list, lists: [...]} 변환."""
    stocks = data.get("stocks", [])
    return {
        "version": 2,
        "active_list": "전체 종목",
        "lists": [
            {
                "name": "전체 종목",
                "config": dict(_DEFAULT_LIST_CONFIG),
                "stocks": stocks,
            }
        ],
    }


def load_watchlist_data():
    """watchlist.json 전체 v2 구조 반환. v1이면 자동 마이그레이션."""
    if not WATCHLIST_PATH.exists():
        return {
            "version": 2,
            "active_list": "전체 종목",
            "lists": [
                {
                    "name": "전체 종목",
                    "config": dict(_DEFAULT_LIST_CONFIG),
                    "stocks": [],
                }
            ],
        }
    try:
        data = json.loads(WATCHLIST_PATH.read_text(encoding="utf-8"))
        if data.get("version") != 2:
            # v1 → v2 마이그레이션
            backup_path = WATCHLIST_PATH.with_suffix(".v1.bak.json")
            backup_path.write_text(
                json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            logger.info(f"watchlist.json v1 백업: {backup_path}")
            data = _migrate_watchlist_v1_to_v2(data)
            save_watchlist_data(data)
            logger.info("watchlist.json v1 → v2 마이그레이션 완료")
        return data
    except Exception as e:
        logger.warning(f"watchlist.json 로드 실패: {e}")
        return {
            "version": 2,
            "active_list": "전체 종목",
            "lists": [
                {
                    "name": "전체 종목",
                    "config": dict(_DEFAULT_LIST_CONFIG),
                    "stocks": [],
                }
            ],
        }


def save_watchlist_data(data):
    """v2 전체 구조를 watchlist.json에 저장."""
    try:
        with open(WATCHLIST_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        total = sum(len(lst["stocks"]) for lst in data.get("lists", []))
        logger.info(f"watchlist.json 저장: {len(data['lists'])}개 리스트, {total}개 종목")
    except Exception as e:
        logger.warning(f"watchlist.json 저장 실패: {e}")


def _find_list(data, list_name):
    """리스트 이름으로 리스트 dict 찾기. 없으면 None."""
    for lst in data.get("lists", []):
        if lst["name"] == list_name:
            return lst
    return None


def _resolve_list_name(data, list_name):
    """list_name이 None이면 active_list 사용."""
    if list_name is None:
        return data.get("active_list", "전체 종목")
    return list_name


def get_list_names():
    """모든 리스트 이름 반환."""
    data = load_watchlist_data()
    return [lst["name"] for lst in data.get("lists", [])]


def get_list(list_name=None):
    """단일 리스트 dict (stocks + config) 반환. 없으면 None."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    return _find_list(data, name)


def create_list(name):
    """빈 리스트 생성. 이미 존재하면 False 반환."""
    data = load_watchlist_data()
    if _find_list(data, name) is not None:
        return False
    data["lists"].append({
        "name": name,
        "config": dict(_DEFAULT_LIST_CONFIG),
        "stocks": [],
    })
    save_watchlist_data(data)
    return True


def delete_list(name):
    """리스트 삭제. 마지막 리스트는 삭제 불가. 성공 시 True."""
    data = load_watchlist_data()
    if len(data["lists"]) <= 1:
        return False
    before = len(data["lists"])
    data["lists"] = [lst for lst in data["lists"] if lst["name"] != name]
    if len(data["lists"]) == before:
        return False
    if data["active_list"] == name:
        data["active_list"] = data["lists"][0]["name"]
    save_watchlist_data(data)
    return True


def rename_list(old_name, new_name):
    """리스트 이름 변경. 성공 시 True."""
    data = load_watchlist_data()
    lst = _find_list(data, old_name)
    if lst is None or _find_list(data, new_name) is not None:
        return False
    lst["name"] = new_name
    if data["active_list"] == old_name:
        data["active_list"] = new_name
    save_watchlist_data(data)
    return True


def set_active_list(name):
    """GUI에서 선택한 활성 리스트 설정."""
    data = load_watchlist_data()
    if _find_list(data, name) is None:
        return False
    data["active_list"] = name
    save_watchlist_data(data)
    return True


def get_list_config(list_name, key):
    """리스트별 config 조회. null이면 글로벌 CONFIG 폴백."""
    lst = get_list(list_name)
    if lst and lst.get("config", {}).get(key) is not None:
        return lst["config"][key]
    return CONFIG.get(key.upper(), None)


def update_list_config(list_name, key, value):
    """리스트별 config 값 설정."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    lst = _find_list(data, name)
    if lst is None:
        return False
    if "config" not in lst:
        lst["config"] = dict(_DEFAULT_LIST_CONFIG)
    lst["config"][key] = value
    save_watchlist_data(data)
    return True


# --- 하위 호환 래퍼 (list_name=None → active_list) ---

def load_watchlist(list_name=None):
    """지정 리스트의 종목 목록 반환."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    lst = _find_list(data, name)
    return lst["stocks"] if lst else []


def save_watchlist(stocks, list_name=None):
    """지정 리스트의 종목 목록 저장."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    lst = _find_list(data, name)
    if lst is None:
        return
    lst["stocks"] = stocks
    save_watchlist_data(data)


def add_to_watchlist(names, list_name=None):
    """종목명 리스트를 워치리스트에 추가 (중복 무시). Returns: 추가된 종목 수."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    lst = _find_list(data, name)
    if lst is None:
        return 0
    existing = {s["name"] for s in lst["stocks"]}
    added = 0
    today = datetime.now(KST).strftime("%Y-%m-%d")
    for n in names:
        if n not in existing:
            lst["stocks"].append({"name": n, "active": True, "added": today})
            existing.add(n)
            added += 1
    if added > 0:
        save_watchlist_data(data)
    return added


def remove_from_watchlist(names, list_name=None):
    """종목명 리스트를 워치리스트에서 삭제. Returns: 삭제된 종목 수."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    lst = _find_list(data, name)
    if lst is None:
        return 0
    names_set = set(names)
    before = len(lst["stocks"])
    lst["stocks"] = [s for s in lst["stocks"] if s["name"] not in names_set]
    removed = before - len(lst["stocks"])
    if removed > 0:
        save_watchlist_data(data)
    return removed


def set_watchlist_active(names, active, list_name=None):
    """지정 종목들의 active 상태 변경."""
    data = load_watchlist_data()
    name = _resolve_list_name(data, list_name)
    lst = _find_list(data, name)
    if lst is None:
        return 0
    names_set = set(names)
    changed = 0
    for s in lst["stocks"]:
        if s["name"] in names_set and s.get("active") != active:
            s["active"] = active
            changed += 1
    if changed > 0:
        save_watchlist_data(data)
    return changed


def get_active_watchlist(list_name=None):
    """active=true인 종목명 리스트만 반환."""
    stocks = load_watchlist(list_name)
    return [s["name"] for s in stocks if s.get("active", True)]


def resolve_ticker(name):
    """
    종목명/티커 → (yfinance 티커, 종목코드, 시장)
    한국: '삼성전자' → ('005930.KS', '005930', 'KOSPI')
    미국: 'AAPL' → ('AAPL', 'AAPL', 'US')
    """
    # 영문 대문자 + 숫자로만 구성되면 US 티커로 간주
    stripped = name.strip().upper()
    if stripped.isascii() and all(c.isalnum() or c in ".-" for c in stripped):
        return stripped, stripped, "US"

    ticker_map = _build_ticker_map()

    if name not in ticker_map:
        # yfinance 검색 폴백
        result = _search_yfinance(name)
        if result:
            return result
        raise ValueError(f"종목을 찾을 수 없습니다: {name}")

    code, market = ticker_map[name]
    suffix = ".KS" if market == "KOSPI" else ".KQ"
    yf_ticker = f"{code}{suffix}"

    return yf_ticker, code, market


def _search_yfinance(name):
    """yfinance로 종목 검색 (폴백)"""
    try:
        import yfinance as yf
        # 한국 시장에서 검색
        results = yf.Search(name, max_results=5)
        for q in results.quotes:
            symbol = q.get("symbol", "")
            if symbol.endswith((".KS", ".KQ")):
                code = symbol.split(".")[0]
                market = "KOSPI" if symbol.endswith(".KS") else "KOSDAQ"
                logger.info(f"yfinance 검색으로 발견: {name} → {symbol}")
                return symbol, code, market
    except Exception as e:
        logger.debug(f"yfinance 검색 실패 ({name}): {e}")
    return None


# ============================================================
# 테마 캐시 (theme_cache.json)
# ============================================================
def _load_theme_cache():
    """theme_cache.json 로드"""
    if not THEME_CACHE_PATH.exists():
        return {}
    try:
        return json.loads(THEME_CACHE_PATH.read_text(encoding="utf-8"))
    except Exception as e:
        logger.warning(f"theme_cache.json 로드 실패: {e}")
        return {}


def _save_theme_cache(cache):
    """theme_cache.json 저장"""
    try:
        with open(THEME_CACHE_PATH, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
        logger.info(f"theme_cache.json 저장: {len(cache)}개 테마")
    except Exception as e:
        logger.warning(f"theme_cache.json 저장 실패: {e}")


def _ask_llm_text(prompt, provider=None):
    """
    텍스트 전용 LLM 호출 (이미지 없음).
    Returns: (응답 텍스트, usage_dict)
    """
    if provider is None:
        provider = LLM_PROVIDER

    if provider == "gemini":
        from google import genai
        api_key = read_secret("gemini_api_key.txt")
        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=[prompt],
        )
        meta = response.usage_metadata
        usage = _make_usage("gemini", GEMINI_MODEL,
                            getattr(meta, "prompt_token_count", 0),
                            getattr(meta, "candidates_token_count", 0))
        return response.text, usage
    elif provider == "claude":
        import anthropic
        api_key = read_secret("anthropic_api_key.txt")
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
        )
        usage = _make_usage("claude", CLAUDE_MODEL,
                            response.usage.input_tokens,
                            response.usage.output_tokens)
        return response.content[0].text, usage
    else:
        raise ValueError(f"지원하지 않는 LLM provider: {provider}")


def refresh_theme_cache(provider=None, log_callback=None):
    """
    LLM에게 테마별 ETF를 조사시키고, yfinance로 검증 후 theme_cache.json에 저장.
    Returns: 갱신된 테마 수
    """
    import yfinance as yf

    if provider is None:
        provider = LLM_PROVIDER

    def log(msg):
        logger.info(msg)
        if log_callback:
            log_callback(msg)

    themes = load_themes()
    cache = _load_theme_cache()
    top_n = ETF_TOP_N
    updated_count = 0

    log(f"{'='*60}")
    log(f"[테마 캐시 갱신] {len(themes)}개 테마 | LLM: {provider}")
    log(f"{'='*60}")

    for t_idx, theme in enumerate(themes, 1):
        theme_name = theme["name"]
        description = theme.get("description", "")

        log(f"\n[{t_idx}/{len(themes)}] {theme_name}")

        # 1단계: LLM에게 ETF 목록 요청
        log(f"  → LLM에게 ETF 목록 조회 중...")
        etf_prompt = (
            f"한국 주식시장에서 '{theme_name}' 테마 관련 ETF를 모두 나열해주세요.\n"
            f"테마 설명: {description}\n\n"
            f"다음 JSON 형식으로만 응답해주세요 (다른 텍스트 없이):\n"
            f'{{"etfs": [{{"name": "ETF 정식명칭", "code": "6자리 종목코드"}}, ...]}}\n\n'
            f"주의: 실제 한국거래소에 상장된 ETF만 포함하세요. "
            f"KODEX, TIGER, ARIRANG, SOL, HANARO, KBSTAR, KOSEF 등 모든 운용사 포함."
        )

        try:
            etf_response, _ = _ask_llm_text(etf_prompt, provider)
            etf_data = _parse_llm_response(etf_response)
            etf_candidates = etf_data.get("etfs", [])
        except Exception as e:
            log(f"  [X] ETF 목록 조회 실패: {e}")
            continue

        if not etf_candidates:
            log(f"  [X] LLM이 ETF를 찾지 못함")
            continue

        log(f"  [후보] {len(etf_candidates)}개 ETF 발견")

        # 2단계: yfinance로 검증 + 거래량 확인
        log(f"  → yfinance 검증 중...")
        valid_etfs = []
        for etf_info in etf_candidates:
            code = str(etf_info.get("code", "")).strip()
            name = etf_info.get("name", "")
            if not code or len(code) != 6 or not code.isdigit():
                continue

            # .KS 먼저, 실패 시 .KQ
            for suffix in [".KS", ".KQ"]:
                yf_ticker = f"{code}{suffix}"
                try:
                    ticker = yf.Ticker(yf_ticker)
                    hist = ticker.history(period="5d")
                    if hist is not None and len(hist) > 0:
                        avg_volume = hist["Volume"].mean()
                        last_close = hist["Close"].iloc[-1]
                        valid_etfs.append({
                            "name": name,
                            "code": code,
                            "yf_ticker": yf_ticker,
                            "avg_volume": float(avg_volume),
                            "last_close": float(last_close),
                            "market_proxy": float(avg_volume * last_close),  # 거래대금 근사치
                        })
                        break
                except Exception:
                    continue

        if not valid_etfs:
            log(f"  [X] 유효한 ETF 없음 (yfinance 검증 실패)")
            continue

        log(f"  [유효] {len(valid_etfs)}개 ETF 검증 통과")

        # 3단계: 거래대금 기준 상위 ETF 선택
        valid_etfs.sort(key=lambda x: x["market_proxy"], reverse=True)
        best_etf = valid_etfs[0]
        log(f"  [선택] {best_etf['name']} ({best_etf['code']})")

        # 4단계: LLM에게 보유종목 조회
        log(f"  → 보유종목 조회 중...")
        holdings_prompt = (
            f"한국 ETF '{best_etf['name']}' (종목코드: {best_etf['code']})의 "
            f"주요 보유종목 상위 {top_n}개를 알려주세요.\n\n"
            f"다음 JSON 형식으로만 응답해주세요 (다른 텍스트 없이):\n"
            f'{{"holdings": ["종목명1", "종목명2", ...]}}\n\n'
            f"주의: 한국 상장 주식의 정확한 한글 종목명을 사용하세요."
        )

        try:
            holdings_response, _ = _ask_llm_text(holdings_prompt, provider)
            holdings_data = _parse_llm_response(holdings_response)
            holdings_raw = holdings_data.get("holdings", [])
        except Exception as e:
            log(f"  [!] 보유종목 조회 실패: {e}, ETF만 캐시")
            holdings_raw = []

        # 5단계: 보유종목을 ticker_cache로 검증
        ticker_map = _build_ticker_map()
        verified_holdings = [h for h in holdings_raw if h in ticker_map]

        if len(verified_holdings) < len(holdings_raw):
            unverified = [h for h in holdings_raw if h not in ticker_map]
            log(f"  [!] 미확인 종목 {len(unverified)}개: {', '.join(unverified[:5])}")

        # 미확인 종목도 포함 (LLM 결과 신뢰)
        final_holdings = holdings_raw[:top_n]
        log(f"  [보유종목] {', '.join(final_holdings[:5])}{'...' if len(final_holdings) > 5 else ''}")

        # 6단계: 캐시 저장
        cache[theme_name] = {
            "etf_code": best_etf["code"],
            "etf_name": best_etf["name"],
            "holdings": final_holdings,
            "updated": datetime.now(KST).strftime("%Y-%m-%d"),
        }
        updated_count += 1
        log(f"  [OK] {theme_name} 캐시 저장 완료")

    _save_theme_cache(cache)

    log(f"\n{'='*60}")
    log(f"[완료] {updated_count}/{len(themes)} 테마 캐시 갱신")
    log(f"{'='*60}")

    return updated_count


# ============================================================
# OHLCV 데이터 수집
# ============================================================
def fetch_ohlcv(ticker_name, period=None):
    """
    종목명으로 OHLCV 데이터 수집.
    yfinance 우선 시도, 실패 시 pykrx 폴백.
    Returns: (pd.DataFrame, yf_ticker, code)
    """
    if period is None:
        period = CHART_PERIOD

    yf_ticker, code, market = resolve_ticker(ticker_name)

    # yfinance 시도
    df = _fetch_yfinance(yf_ticker, period)
    if df is not None and len(df) > 0:
        logger.info(f"[yfinance] {ticker_name} ({yf_ticker}): {len(df)}개 데이터")
        return df, yf_ticker, code

    # pykrx 폴백
    logger.warning(f"[yfinance] {ticker_name} 실패, pykrx 폴백 시도")
    df = _fetch_pykrx(code, period)
    if df is not None and len(df) > 0:
        logger.info(f"[pykrx] {ticker_name} ({code}): {len(df)}개 데이터")
        return df, yf_ticker, code

    raise RuntimeError(f"{ticker_name} 데이터를 가져올 수 없습니다 (yfinance, pykrx 모두 실패)")


def _fetch_yfinance(yf_ticker, period):
    """yfinance로 OHLCV 데이터 수집"""
    try:
        import yfinance as yf
        ticker = yf.Ticker(yf_ticker)
        df = ticker.history(period=period)
        if df.empty:
            return None
        # 컬럼 통일
        df = df[["Open", "High", "Low", "Close", "Volume"]].copy()
        df.index.name = "Date"
        return df
    except Exception as e:
        logger.debug(f"yfinance 오류 ({yf_ticker}): {e}")
        return None


def _fetch_pykrx(code, period):
    """pykrx로 OHLCV 데이터 수집"""
    try:
        from pykrx import stock

        # period 문자열을 날짜 범위로 변환
        end_date = datetime.now(KST).strftime("%Y%m%d")
        period_map = {"1mo": 30, "3mo": 90, "6mo": 180, "1y": 365, "2y": 730}
        days = period_map.get(period, 365)
        start_date = (datetime.now(KST) - timedelta(days=days)).strftime("%Y%m%d")

        df = stock.get_market_ohlcv_by_date(start_date, end_date, code)
        if df.empty:
            return None

        # 컬럼명 통일 (한글 → 영문)
        col_map = {"시가": "Open", "고가": "High", "저가": "Low", "종가": "Close", "거래량": "Volume"}
        df = df.rename(columns=col_map)
        df = df[["Open", "High", "Low", "Close", "Volume"]].copy()
        df.index.name = "Date"
        return df
    except Exception as e:
        logger.debug(f"pykrx 오류 ({code}): {e}")
        return None


# ============================================================
# 캔들차트 생성
# ============================================================
def generate_chart(ticker_name, df, code=None, market=None):
    """
    mplfinance로 캔들스틱 차트 생성 (MA + 거래량 포함).
    Returns: 저장된 이미지 경로 (Path)
    """
    import mplfinance as mpf
    import matplotlib
    matplotlib.use("Agg")  # headless 호환

    # 한글 폰트 설정
    import matplotlib.font_manager as fm
    kr_font = None
    for font_name in ["Malgun Gothic", "NanumGothic", "AppleGothic"]:
        if any(font_name in f.name for f in fm.fontManager.ttflist):
            kr_font = font_name
            break
    if kr_font:
        matplotlib.rcParams["font.family"] = kr_font
        matplotlib.rcParams["axes.unicode_minus"] = False

    title_str = f"{ticker_name}"
    if code:
        title_str += f" ({code})"

    # MA 색상: 5일(빨강), 20일(초록), 60일(파랑), 120일(보라)
    ma_colors = ["#FF6B6B", "#4ECDC4", "#45B7D1", "#96CEB4"]

    # 유효한 MA만 사용 (데이터 길이보다 긴 MA는 제외)
    valid_ma = [m for m in CHART_MA_LINES if m < len(df)]
    valid_colors = ma_colors[:len(valid_ma)]

    # 차트 스타일: 시장에 따라 색상 전환
    if market == "US":
        up_color, down_color = "#26A69A", "#EF5350"  # 미국식: 초록/빨강
    else:
        up_color, down_color = "#FF3B30", "#007AFF"  # 한국식: 빨강/파랑
    mc = mpf.make_marketcolors(
        up=up_color, down=down_color,
        edge="inherit",
        wick="inherit",
        volume="in",
    )
    style = mpf.make_mpf_style(
        marketcolors=mc,
        figcolor="white",
        gridcolor="#F0F0F0",
        gridstyle="-",
        y_on_right=True,
        rc={"font.size": 10, "font.family": kr_font or "sans-serif", "axes.unicode_minus": False},
    )

    # 파일명 생성
    date_str = datetime.now(KST).strftime("%Y%m%d")
    filename = f"{ticker_name}_{date_str}.png"
    save_path = CHARTS_DIR / filename

    kwargs = dict(
        type="candle",
        volume=True,
        style=style,
        figsize=(12, 8),
        title=title_str,
        savefig=dict(fname=str(save_path), dpi=150, bbox_inches="tight"),
    )

    if valid_ma:
        kwargs["mav"] = tuple(valid_ma)
        kwargs["mavcolors"] = valid_colors

    mpf.plot(df, **kwargs)

    logger.info(f"차트 생성: {save_path}")
    return save_path


# ============================================================
# LLM 차트 분석
# ============================================================
def load_chart_prompt():
    """chart_prompt.txt 로드"""
    prompt_path = BASE_DIR / "chart_prompt.txt"
    if not prompt_path.exists():
        raise FileNotFoundError(f"프롬프트 파일을 찾을 수 없습니다: {prompt_path}")
    return prompt_path.read_text(encoding="utf-8")


def analyze_chart_with_llm(chart_image_path, ticker_name, provider=None):
    """
    차트 이미지를 LLM 비전 모델에 보내 기술적 분석 수행.
    Returns: 분석 결과 dict (grade, confidence, trend, signals, ...)
    """
    if provider is None:
        provider = LLM_PROVIDER

    # 프롬프트 준비
    prompt_template = load_chart_prompt()
    prompt = prompt_template.replace("{ticker_name}", ticker_name)

    # 이미지 읽기
    image_bytes = Path(chart_image_path).read_bytes()
    image_b64 = base64.standard_b64encode(image_bytes).decode("utf-8")

    # LLM 호출
    if provider == "claude":
        raw_response, usage = _analyze_with_claude(image_b64, prompt)
    elif provider == "gemini":
        raw_response, usage = _analyze_with_gemini(image_b64, prompt)
    else:
        raise ValueError(f"지원하지 않는 LLM provider: {provider}")

    # JSON 파싱
    result = _parse_llm_response(raw_response)
    result["ticker_name"] = ticker_name
    result["token_usage"] = usage
    return result


def _analyze_with_claude(image_b64, prompt):
    """Claude 비전 API로 차트 분석. Returns: (text, usage_dict)"""
    import anthropic

    api_key = read_secret("anthropic_api_key.txt")
    client = anthropic.Anthropic(api_key=api_key)

    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": "image/png",
                        "data": image_b64,
                    },
                },
                {
                    "type": "text",
                    "text": prompt,
                },
            ],
        }],
    )

    usage = _make_usage("claude", CLAUDE_MODEL,
                        response.usage.input_tokens,
                        response.usage.output_tokens)
    return response.content[0].text, usage


def _analyze_with_gemini(image_b64, prompt):
    """Gemini 비전 API로 차트 분석 (google.genai SDK). Returns: (text, usage_dict)"""
    from google import genai
    from google.genai import types

    api_key = read_secret("gemini_api_key.txt")
    client = genai.Client(api_key=api_key)

    # base64 이미지를 Part로 변환
    image_bytes = base64.standard_b64decode(image_b64)
    image_part = types.Part.from_bytes(data=image_bytes, mime_type="image/png")

    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=[image_part, prompt],
    )

    meta = response.usage_metadata
    usage = _make_usage("gemini", GEMINI_MODEL,
                        getattr(meta, "prompt_token_count", 0),
                        getattr(meta, "candidates_token_count", 0))
    return response.text, usage


def _parse_llm_response(raw_text):
    """LLM 응답에서 텍스트 기반 분석 결과 파싱 (v2 프롬프트)"""
    result = {"raw_response": raw_text, "parse_error": False}

    # 1. 결론 (매력도): "결론: A-1", "결론: **D (리스크)**", "## 결론: **A-2**" 등
    m = re.search(r"결론\s*[:：]\s*\*{0,2}\s*(?:매력도\s*)?([A-Da-d]-?[12]?)\b", raw_text)
    if m:
        grade_raw = m.group(1).upper()
        # "A1" → "A-1", "A2" → "A-2" 정규화
        if grade_raw in ("A1", "A2"):
            grade_raw = grade_raw[0] + "-" + grade_raw[1]
        result["grade"] = grade_raw
    else:
        result["grade"] = "N/A"
        result["parse_error"] = True
        logger.error(f"등급 파싱 실패. 원본 응답:\n{raw_text[:500]}")

    # 2. 결론 일치 횟수: "3회 중 2회"
    m = re.search(r"결론\s*일치\s*횟수\s*[:：]\s*(.+)", raw_text)
    result["consensus_count"] = m.group(1).strip() if m else ""

    # 3. 신뢰도 등급: High/Medium/Low (한글도 대응)
    m = re.search(r"신뢰도\s*등급\s*[:：]\s*(High|Medium|Low|높음|보통|낮음)", raw_text, re.IGNORECASE)
    if m:
        rel = m.group(1).capitalize()
        # 한글 → 영문 변환
        rel_map = {"높음": "High", "보통": "Medium", "낮음": "Low"}
        result["reliability"] = rel_map.get(rel, rel)
    else:
        result["reliability"] = ""

    # 4. 핵심 근거 (결론 블록 이후 ~ 현재가 or 끝)
    m = re.search(
        r"핵심\s*근거\s*[:：]?\s*\n?(.*?)(?=\n\s*\d+\.\s*현재가|\n\s*현재가\s*[:：]|\Z)",
        raw_text, re.DOTALL
    )
    result["reasoning"] = m.group(1).strip() if m else ""

    # 5. A-1/A-2 전용 필드
    if result["grade"] in ("A-1", "A-2"):
        m = re.search(r"현재가\s*[:：]\s*(.+)", raw_text)
        result["current_price"] = m.group(1).strip() if m else ""

        m = re.search(r"목표가\s*[:：]\s*(.+)", raw_text)
        result["target_price"] = m.group(1).strip() if m else ""

        m = re.search(r"매수\s*전략\s*[:：]\s*(.+)", raw_text)
        result["buy_strategy"] = m.group(1).strip() if m else ""

        m = re.search(r"매도\s*전략\s*[:：]\s*(.+)", raw_text)
        result["sell_strategy"] = m.group(1).strip() if m else ""

    # reliability → confidence 숫자 변환 (기존 정렬 로직 호환)
    confidence_map = {"High": 90, "Medium": 70, "Low": 50}
    result["confidence"] = confidence_map.get(result.get("reliability", ""), 0)

    return result


# ============================================================
# 분석 파이프라인 (Feature 1: 종목 분석)
# ============================================================
def run_stock_analysis(ticker_names, provider=None, log_callback=None):
    """
    종목 리스트에 대해 차트 기반 기술적 분석 수행.
    Returns: {results: list, a_rated: list, summary: str}
    """
    if provider is None:
        provider = LLM_PROVIDER

    def log(msg):
        logger.info(msg)
        if log_callback:
            log_callback(msg)

    results = []
    errors = []
    total = len(ticker_names)
    total_usage = {
        "input_tokens": 0, "output_tokens": 0, "total_tokens": 0,
        "input_cost_usd": 0.0, "output_cost_usd": 0.0, "total_cost_usd": 0.0,
        "api_calls": 0,
    }

    log(f"{'='*60}")
    log(f"[KR Chart Rater] 종목 분석 시작")
    log(f"종목 수: {total}개 | LLM: {provider}")
    log(f"{'='*60}")

    for i, name in enumerate(ticker_names, 1):
        log(f"\n[{i}/{total}] {name}")

        try:
            # 1. 티커 해석
            log(f"  → 티커 조회 중...")
            yf_ticker, code, market = resolve_ticker(name)
            log(f"  [티커] {yf_ticker} ({market})")

            # 2. 데이터 수집
            log(f"  → 데이터 수집 중...")
            df, _, _ = fetch_ohlcv(name)
            log(f"  [데이터] {len(df)}개 봉 ({df.index[0].strftime('%Y-%m-%d')} ~ {df.index[-1].strftime('%Y-%m-%d')})")

            # 3. 차트 생성
            log(f"  → 차트 생성 중...")
            chart_path = generate_chart(name, df, code)
            log(f"  [차트] {chart_path.name}")

            # 4. LLM 분석
            log(f"  → LLM 분석 중 ({provider})...")
            start_time = time.time()
            analysis = analyze_chart_with_llm(chart_path, name, provider)
            elapsed = time.time() - start_time
            usage = analysis.get("token_usage", {})
            if usage:
                cost_usd = usage["total_cost_usd"]
                cost_krw = cost_usd * 1400
                cost_str = f" | 토큰: {usage.get('total_tokens', 0):,} | ${cost_usd:.4f} ({cost_krw:.0f}원)"
            else:
                cost_str = ""
            log(f"  [분석] 등급: {analysis.get('grade', 'N/A')} | 신뢰도: {analysis.get('confidence', 0)}% ({elapsed:.1f}s){cost_str}")
            if usage:
                _accumulate_usage(total_usage, usage)
                cum_usd = total_usage["total_cost_usd"]
                cum_krw = cum_usd * 1400
                log(f"  [누적 비용] ${cum_usd:.4f} ({cum_krw:.0f}원)")

            # 파싱 실패 시 1회 재시도
            if analysis.get("parse_error"):
                log(f"  [!] JSON 파싱 실패, 재시도 중...")
                analysis = analyze_chart_with_llm(chart_path, name, provider)
                retry_usage = analysis.get("token_usage", {})
                if retry_usage:
                    _accumulate_usage(total_usage, retry_usage)
                if analysis.get("parse_error"):
                    log(f"  [X] 재시도도 실패")

            analysis["chart_path"] = str(chart_path)
            analysis["yf_ticker"] = yf_ticker
            analysis["code"] = code
            analysis["market"] = market
            results.append(analysis)
            log(f"  [OK] {name} 완료")

        except Exception as e:
            error_msg = f"{name}: {e}"
            logger.error(error_msg, exc_info=True)
            log(f"  [X] 오류: {e}")
            errors.append(error_msg)

    # A등급 필터 및 정렬 (A-1 우선, A-2 다음, confidence 내림차순)
    a_rated = [r for r in results if r.get("grade", "").startswith("A")]
    a_rated.sort(key=lambda x: (0 if x.get("grade") == "A-1" else 1, -x.get("confidence", 0)))

    # 요약
    log(f"\n{'='*60}")
    log(f"[완료] {len(results)}/{total} 분석 | A-1/A-2 선정: {len(a_rated)}개 | 오류: {len(errors)}개")
    if total_usage["api_calls"] > 0:
        cost_usd = total_usage["total_cost_usd"]
        cost_krw = cost_usd * 1400  # 근사 환율
        log(f"[토큰] 입력: {total_usage['input_tokens']:,} | 출력: {total_usage['output_tokens']:,} | "
            f"합계: {total_usage['total_tokens']:,}")
        log(f"[비용] ${cost_usd:.4f} (약 {cost_krw:.0f}원) | API 호출: {total_usage['api_calls']}회")
    log(f"{'='*60}")

    if a_rated:
        log(f"\n[선정] A-1/A-2 종목 (매력도순):")
        for idx, r in enumerate(a_rated, 1):
            log(f"  {idx}. [{r.get('grade', '')}] {r.get('ticker_name', r.get('ticker', ''))} "
                f"(신뢰도 {r.get('reliability', '')}) - "
                f"{r.get('reasoning', '')[:60]}...")

    # 결과 저장
    output = {
        "analysis_date": datetime.now(KST).isoformat(),
        "provider": provider,
        "total_analyzed": len(results),
        "total_errors": len(errors),
        "results": results,
        "a_rated": a_rated,
        "errors": errors,
        "token_usage": total_usage,
    }

    result_path = save_results_json(output, "stocks")
    log(f"\n결과 저장: {result_path}")

    # 분석 성공한 종목을 워치리스트에 자동 추가
    analyzed_names = [r.get("ticker_name", "") for r in results if r.get("ticker_name")]
    if analyzed_names:
        added = add_to_watchlist(analyzed_names)
        if added > 0:
            log(f"[워치리스트] {added}개 종목 자동 추가")

    return output


# ============================================================
# 테마 분석 (Feature 2)
# ============================================================
def load_themes():
    """themes.json에서 테마 목록 로드"""
    themes_path = BASE_DIR / "themes.json"
    if not themes_path.exists():
        raise FileNotFoundError(f"테마 파일을 찾을 수 없습니다: {themes_path}")
    return json.loads(themes_path.read_text(encoding="utf-8"))


def find_etf_for_theme(theme_name):
    """
    theme_cache.json에서 테마의 대표 ETF 정보 조회.
    Returns: {etf_name, ticker_code} 또는 None
    """
    cache = _load_theme_cache()
    entry = cache.get(theme_name)
    if not entry:
        logger.warning(f"테마 '{theme_name}' 캐시 없음. 'refresh-themes' 명령으로 캐시를 갱신하세요.")
        return None

    result = {
        "etf_name": entry.get("etf_name", ""),
        "ticker_code": entry.get("etf_code", ""),
    }
    logger.info(f"테마 '{theme_name}' → ETF: {result['etf_name']} ({result['ticker_code']})")
    return result


def get_etf_holdings(theme_name, top_n=None):
    """
    theme_cache.json에서 테마의 ETF 보유종목 반환.
    Returns: [{"name": str}, ...]
    """
    if top_n is None:
        top_n = ETF_TOP_N

    cache = _load_theme_cache()
    entry = cache.get(theme_name)
    if not entry:
        logger.warning(f"테마 '{theme_name}' 캐시 없음")
        return []

    holdings_names = entry.get("holdings", [])[:top_n]
    holdings = [{"name": name} for name in holdings_names]

    logger.info(f"테마 '{theme_name}' 보유종목: {len(holdings)}개 (상위 {top_n})")
    return holdings


def run_theme_analysis(theme_names=None, top_n=None, provider=None, log_callback=None):
    """
    테마별 ETF 보유종목에 대해 차트 분석 수행.
    theme_names가 None이면 전체 테마 분석.
    Returns: {themes: list[theme_result]}
    """
    if provider is None:
        provider = LLM_PROVIDER
    if top_n is None:
        top_n = ETF_TOP_N

    def log(msg):
        logger.info(msg)
        if log_callback:
            log_callback(msg)

    all_themes = load_themes()

    # 선택된 테마만 필터
    if theme_names:
        themes = [t for t in all_themes if t["name"] in theme_names]
        if not themes:
            log(f"[!] 선택된 테마를 찾을 수 없습니다: {theme_names}")
            return {"themes": []}
    else:
        themes = all_themes

    log(f"{'='*60}")
    log(f"[KR Chart Rater] 테마 분석 시작")
    log(f"테마 수: {len(themes)}개 | 보유종목 Top: {top_n} | LLM: {provider}")
    log(f"{'='*60}")

    theme_results = []

    for t_idx, theme in enumerate(themes, 1):
        theme_name = theme["name"]
        log(f"\n{'-'*40}")
        log(f"[테마 {t_idx}/{len(themes)}] {theme_name}")

        # 1. ETF 탐색 (캐시에서)
        log(f"  → ETF 조회 중...")
        etf = find_etf_for_theme(theme_name)
        if not etf:
            log(f"  [X] 캐시에 ETF 정보 없음. 'refresh-themes' 명령을 먼저 실행하세요.")
            theme_results.append({
                "theme": theme_name,
                "etf": None,
                "error": "캐시에 ETF 없음 (refresh-themes 필요)",
                "holdings_analyzed": [],
                "a_rated": [],
            })
            continue

        log(f"  [ETF] {etf['etf_name']} ({etf['ticker_code']})")

        # 2. 보유종목 추출 (캐시에서)
        log(f"  → 보유종목 조회 중 (Top {top_n})...")
        holdings = get_etf_holdings(theme_name, top_n)
        if not holdings:
            log(f"  [X] 보유종목 조회 실패")
            theme_results.append({
                "theme": theme_name,
                "etf": etf,
                "error": "보유종목 조회 실패",
                "holdings_analyzed": [],
                "a_rated": [],
            })
            continue

        holding_names = [h["name"] for h in holdings]
        log(f"  [보유종목] {', '.join(holding_names[:5])}{'...' if len(holding_names) > 5 else ''}")

        # 3. 보유종목 차트 분석
        log(f"  → 보유종목 차트 분석 시작...")
        analysis = run_stock_analysis(holding_names, provider, log_callback)

        theme_results.append({
            "theme": theme_name,
            "etf": etf,
            "holdings": holdings,
            "holdings_analyzed": analysis.get("results", []),
            "a_rated": analysis.get("a_rated", []),
            "errors": analysis.get("errors", []),
        })

    # 전체 요약
    total_a = sum(len(t.get("a_rated", [])) for t in theme_results)
    log(f"\n{'='*60}")
    log(f"[테마 분석 완료] {len(theme_results)}개 테마 | 전체 A-1/A-2 선정: {total_a}개")
    log(f"{'='*60}")

    for tr in theme_results:
        a_list = tr.get("a_rated", [])
        if a_list:
            names = ", ".join(r.get("ticker_name", "") for r in a_list)
            log(f"  {tr['theme']}: A-1/A-2 {len(a_list)}개 ({names})")

    # 결과 저장
    output = {
        "analysis_date": datetime.now(KST).isoformat(),
        "provider": provider,
        "themes": theme_results,
    }

    result_path = save_results_json(output, "themes")
    log(f"\n결과 저장: {result_path}")

    return output


# ============================================================
# 결과 저장
# ============================================================
def save_results_json(data, prefix="analysis"):
    """결과를 JSON 파일로 저장"""
    timestamp = datetime.now(KST).strftime("%Y%m%d_%H%M%S")
    filename = f"{prefix}_{timestamp}.json"
    path = RESULTS_DIR / filename

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2, default=str)

    logger.info(f"결과 저장: {path}")
    return path


# ============================================================
# Word 보고서 내보내기
# ============================================================
# 등급별 색상 (RGB 튜플)
_GRADE_COLORS = {
    "A-1": (220, 38, 38),    # 빨강 (속도형 매력)
    "A-2": (234, 88, 12),    # 주황 (완만추세형)
    "B":   (202, 138, 4),    # 노랑
    "C":   (107, 114, 128),  # 회색
    "D":   (128, 128, 128),  # 진회색
    "A":   (220, 38, 38),    # 레거시 호환
    "N/A": (107, 114, 128),
}


def save_results_docx(data, prefix="analysis"):
    """
    분석 결과 dict → Word 보고서 (.docx)
    data: run_stock_analysis() 또는 run_theme_analysis()의 반환값
    Returns: 저장된 파일 경로 (Path)
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # ── 페이지 여백 ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── 기본 스타일 ──
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # 헤딩 스타일 (NP Orange 계열)
    _HCFG = {
        1: (18, (200, 90, 40)),     # 큰 제목 - 주황갈색
        2: (14, (61, 43, 31)),      # 중 제목 - TEXT_DARK
        3: (12, (100, 70, 50)),     # 소 제목
    }
    for lvl, (sz, rgb) in _HCFG.items():
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "맑은 고딕"
        hs.font.size = Pt(sz)
        hs.font.color.rgb = RGBColor(*rgb)
        hs.font.bold = True
        hs.paragraph_format.space_before = Pt(max(6, 24 - lvl * 4))
        hs.paragraph_format.space_after = Pt(4)

    # ── 헬퍼 함수 ──
    def _set_cell_shading(cell, color_hex):
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex})
        tc_pr.append(shd)

    def _add_hr():
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bot = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "1", qn("w:color"): "F0D0B0"})
        pBdr.append(bot)
        pPr.append(pBdr)

    def _grade_run(para, grade):
        """등급 텍스트를 색상 run으로 추가"""
        r = para.add_run(f" [{grade}] ")
        rgb = _GRADE_COLORS.get(grade, (107, 114, 128))
        r.font.color.rgb = RGBColor(*rgb)
        r.bold = True
        r.font.size = Pt(14)
        return r

    # ── 제목 페이지 ──
    analysis_date = data.get("analysis_date", datetime.now(KST).isoformat())
    if isinstance(analysis_date, str):
        try:
            dt = datetime.fromisoformat(analysis_date)
        except Exception:
            dt = datetime.now(KST)
    else:
        dt = analysis_date
    date_str = dt.strftime("%Y년 %m월 %d일")

    is_theme = "themes" in data
    title_text = "테마별 차트 분석 보고서" if is_theme else "종목 차트 분석 보고서"

    h = doc.add_heading(title_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 제목 아래 밑줄
    pPr = h._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bot = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "12",
        qn("w:space"): "4", qn("w:color"): "E8864A"})
    pBdr.append(bot)
    pPr.append(pBdr)

    # 분석 정보
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(12)
    r = info_p.add_run(f"분석일: {date_str}  |  LLM: {data.get('provider', 'N/A').upper()}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    # 토큰 비용 표시
    usage = data.get("token_usage")
    if usage and usage.get("api_calls", 0) > 0:
        cost_p = doc.add_paragraph()
        cost_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cost_usd = usage["total_cost_usd"]
        cost_krw = cost_usd * 1400
        cr = cost_p.add_run(
            f"토큰: {usage['total_tokens']:,}  |  "
            f"비용: ${cost_usd:.4f} (약 {cost_krw:.0f}원)  |  "
            f"API 호출: {usage['api_calls']}회"
        )
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(130, 130, 130)

    # ── 종목 분석 보고서 ──
    if not is_theme:
        results = data.get("results", [])
        a_rated = data.get("a_rated", [])

        # 요약 테이블
        doc.add_heading("분석 요약", level=2)
        total = data.get("total_analyzed", len(results))
        errors = data.get("total_errors", 0)
        p = doc.add_paragraph()
        p.add_run(f"분석 종목: {total}개  |  오류: {errors}개  |  A-1/A-2 선정: {len(a_rated)}개")

        if results:
            _write_summary_table(doc, results, _set_cell_shading, _grade_run)

        _add_hr()

        # 개별 종목 상세
        doc.add_heading("종목별 상세 분석", level=2)
        for idx, r in enumerate(results):
            _write_stock_detail(doc, r, idx + 1, _set_cell_shading, _grade_run, _add_hr)

    # ── 테마 분석 보고서 ──
    else:
        theme_results = data.get("themes", [])
        total_a = sum(len(t.get("a_rated", [])) for t in theme_results)

        doc.add_heading("분석 요약", level=2)
        p = doc.add_paragraph()
        p.add_run(f"테마 수: {len(theme_results)}개  |  전체 A-1/A-2 선정: {total_a}개")

        for t_idx, theme_data in enumerate(theme_results, 1):
            theme_name = theme_data.get("theme", f"테마 {t_idx}")
            etf = theme_data.get("etf")
            results = theme_data.get("holdings_analyzed", [])
            a_rated = theme_data.get("a_rated", [])
            error_msg = theme_data.get("error")

            _add_hr()
            doc.add_heading(f"{t_idx}. {theme_name}", level=2)

            if etf:
                p = doc.add_paragraph()
                p.add_run("ETF: ").bold = True
                p.add_run(f"{etf.get('etf_name', '')} ({etf.get('ticker_code', '')})")

            if error_msg:
                p = doc.add_paragraph()
                r = p.add_run(f"[오류] {error_msg}")
                r.font.color.rgb = RGBColor(220, 38, 38)
                continue

            if results:
                _write_summary_table(doc, results, _set_cell_shading, _grade_run)

                doc.add_heading(f"{theme_name} - 종목별 상세", level=3)
                for idx, r in enumerate(results):
                    _write_stock_detail(doc, r, idx + 1, _set_cell_shading, _grade_run, _add_hr)

    # ── 푸터 ──
    _add_hr()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("KR Chart Rater - AI 기반 차트 기술적 분석 보고서")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(150, 150, 150)

    # ── 저장 ──
    timestamp = datetime.now(KST).strftime("%Y%m%d_%H%M%S")
    filename = f"{prefix}_{timestamp}.docx"
    path = RESULTS_DIR / filename
    doc.save(str(path))
    logger.info(f"Word 보고서 저장: {path}")
    return path


def _write_summary_table(doc, results, _set_cell_shading, _grade_run):
    """종목 요약 테이블 생성"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    headers = ["종목명", "등급", "신뢰도", "추세", "MA 배열"]
    tbl = doc.add_table(rows=1 + len(results), cols=len(headers))
    tbl.style = "Table Grid"

    # 헤더 행
    for ci, text in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, "C2652E")  # ACCENT_DARK

    # 데이터 행
    for ri, result in enumerate(results):
        grade = result.get("grade", "N/A")
        trend = result.get("trend", {})
        row = tbl.rows[ri + 1]

        vals = [
            result.get("ticker_name", result.get("ticker", "")),
            grade,
            f"{result.get('confidence', 0)}%",
            trend.get("direction", "-"),
            trend.get("ma_arrangement", "-"),
        ]
        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)

            # 등급 컬럼 색상
            if ci == 1:
                rgb = _GRADE_COLORS.get(grade, (107, 114, 128))
                r.font.color.rgb = RGBColor(*rgb)
                r.bold = True

            # 짝수 행 배경
            if ri % 2 == 1:
                _set_cell_shading(cell, "FFF3EB")  # ACCENT_LIGHT

    doc.add_paragraph()  # 여백


def _write_stock_detail(doc, result, num, _set_cell_shading, _grade_run, _add_hr):
    """개별 종목 상세 분석 작성"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    name = result.get("ticker_name", result.get("ticker", f"종목 {num}"))
    grade = result.get("grade", "N/A")
    confidence = result.get("confidence", 0)
    trend = result.get("trend", {})

    # 종목 제목 + 등급
    h = doc.add_heading(level=3)
    h.add_run(f"{num}. {name}")
    _grade_run(h, grade)
    r = h.add_run(f"  (신뢰도 {confidence}%)")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)
    r.bold = False

    # 기본 정보
    code = result.get("code", "")
    market = result.get("market", "")
    if code:
        p = doc.add_paragraph()
        p.add_run("종목코드: ").bold = True
        p.add_run(f"{code} ({market})")
        p.add_run("  |  추세: ").bold = True
        p.add_run(f"{trend.get('direction', '-')} ({trend.get('ma_arrangement', '-')}, 강도: {trend.get('strength', '-')})")

    # 차트 이미지 삽입
    chart_path = result.get("chart_path")
    if chart_path and Path(chart_path).exists():
        try:
            doc.add_picture(str(chart_path), width=Inches(5.5))
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            p = doc.add_paragraph()
            r = p.add_run(f"(차트 이미지 삽입 실패: {e})")
            r.font.color.rgb = RGBColor(150, 150, 150)
            r.font.size = Pt(8)

    # 기술적 신호
    signals = result.get("signals", [])
    if signals:
        p = doc.add_paragraph()
        r = p.add_run("기술적 신호")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(30, 64, 175)
        for sig in signals:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(sig).font.size = Pt(9)

    # 리스크 요인
    risk_factors = result.get("risk_factors", [])
    if risk_factors:
        p = doc.add_paragraph()
        r = p.add_run("리스크 요인")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(220, 38, 38)
        for rf in risk_factors:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(rf).font.size = Pt(9)

    # 종합 의견
    reasoning = result.get("reasoning", "")
    if reasoning:
        p = doc.add_paragraph()
        r = p.add_run("종합 의견")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(61, 43, 31)
        p = doc.add_paragraph()
        p.add_run(reasoning).font.size = Pt(9.5)

    # 목표가 / 손절가
    target = result.get("price_target_zone", "")
    stoploss = result.get("stop_loss_zone", "")
    if target or stoploss:
        p = doc.add_paragraph()
        r = p.add_run("매매 전략")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(61, 43, 31)
        if target:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run("목표가: ").bold = True
            bp.add_run(target).font.size = Pt(9)
        if stoploss:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run("손절가: ").bold = True
            bp.add_run(stoploss).font.size = Pt(9)

    _add_hr()


# ============================================================
# 이메일 보고서 발송
# ============================================================
def build_email_body(result, provider):
    """분석 결과로 이메일 본문 생성. GUI와 스케줄 양쪽에서 재사용."""
    dt = datetime.now(KST)
    date_str = dt.strftime("%Y년 %m월 %d일")
    total = result.get("total_analyzed", 0)
    errors = result.get("total_errors", 0)
    a_rated = result.get("a_rated", [])

    lines = [
        "KR Chart Rater 자동 분석 보고서",
        f"분석일: {date_str}",
        f"LLM: {provider.upper()}",
        "",
        f"분석 종목: {total}개 | 오류: {errors}개 | A-1/A-2 선정: {len(a_rated)}개",
        "",
    ]

    if a_rated:
        lines.append("A-1/A-2 선정 종목:")
        for idx, r in enumerate(a_rated, 1):
            name = r.get("ticker_name", "")
            grade = r.get("grade", "")
            reliability = r.get("reliability", "")
            lines.append(f"  {idx}. [{grade}] {name} (신뢰도: {reliability})")
    else:
        lines.append("A-1/A-2 선정 종목 없음")

    usage = result.get("token_usage", {})
    if usage.get("api_calls", 0) > 0:
        cost_usd = usage["total_cost_usd"]
        cost_krw = cost_usd * 1400
        lines.extend([
            "",
            f"토큰 사용: {usage['total_tokens']:,}개 "
            f"(입력 {usage['input_tokens']:,} / 출력 {usage['output_tokens']:,})",
            f"비용: ${cost_usd:.4f} (약 {cost_krw:.0f}원) | API 호출 {usage['api_calls']}회",
        ])

    lines.extend([
        "",
        "상세 분석은 첨부된 Word 보고서를 참고하세요.",
        "",
        "- KR Chart Rater",
    ])

    subject = f"[KR Chart Rater] {date_str} 차트 분석 보고서"
    return subject, "\n".join(lines)


def send_report_email(to_addr, subject, body_text, attachment_path=None):
    """
    SMTP로 보고서 이메일 발송.
    config.txt의 EMAIL_* 설정 사용.
    """
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.application import MIMEApplication

    smtp_host = CONFIG.get("EMAIL_SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(CONFIG.get("EMAIL_SMTP_PORT", "587"))
    from_addr = CONFIG.get("EMAIL_FROM", "")
    password_file = CONFIG.get("EMAIL_PASSWORD_FILE", "email_password.txt")

    if not from_addr:
        raise ValueError("config.txt에 EMAIL_FROM이 설정되지 않았습니다.")

    password = read_secret(password_file)

    msg = MIMEMultipart()
    msg["From"] = from_addr
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.attach(MIMEText(body_text, "plain", "utf-8"))

    if attachment_path:
        path = Path(attachment_path)
        if path.exists():
            with open(path, "rb") as f:
                part = MIMEApplication(f.read(), Name=path.name)
            part["Content-Disposition"] = f'attachment; filename="{path.name}"'
            msg.attach(part)

    with smtplib.SMTP(smtp_host, smtp_port) as server:
        server.starttls()
        server.login(from_addr, password)
        server.send_message(msg)

    logger.info(f"이메일 발송 완료: {to_addr} ({subject})")
